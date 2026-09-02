import json
import os
import re
from typing import Optional

import requests
from config import Config

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_OK = True
except ImportError:
    SKLEARN_OK = False


# Groq model compatibility (Groq retired Llama 3.1 8B on 2026-08-16).
# We resolve the active model from the account when possible so stale .env files
# cannot keep the app pinned to a retired model.
_MODEL_CACHE = {"model": "openai/gpt-oss-20b", "expires": 0.0}

# Never select the 120B model for this app: the user's current org limit is 8K TPM,
# and stale .env values were previously overriding the intended 20B model.
_PRIMARY_MODEL = "openai/gpt-oss-20b"
_WEB_MODEL = "groq/compound-mini"
_BLOCKED_MODELS = {"openai/gpt-oss-120b", "llama-3.1-8b-instant"}

def _resolve_llm_model(force_refresh=False):
    import time
    now = time.time()
    if not force_refresh and _MODEL_CACHE["expires"] > now:
        return _MODEL_CACHE["model"]
    chosen = _PRIMARY_MODEL
    if Config.LLM_API_KEY:
        try:
            r = requests.get(
                Config.GROQ_BASE_URL.rstrip("/") + "/models",
                headers={"Authorization": f"Bearer {Config.LLM_API_KEY}"},
                timeout=5,
            )
            if r.ok:
                ids = {str(x.get("id")) for x in (r.json().get("data") or []) if isinstance(x, dict) and x.get("id")}
                if _PRIMARY_MODEL not in ids:
                    # Pick a small compatible text model, never 120B.
                    for candidate in ("qwen/qwen3.8-27b", "qwen/qwen3.6-27b"):
                        if candidate in ids:
                            chosen = candidate
                            break
        except Exception:
            pass
    _MODEL_CACHE.update({"model": chosen, "expires": now + 300})
    return chosen

def _is_compound_model(model=None):
    return (model or Config.WEB_MODEL) in {"groq/compound", "groq/compound-mini"}


# ----------------------------------------------------------------------
# Document extraction
# ----------------------------------------------------------------------
def extract_text(filepath, ext):
    ext = ext.lower().lstrip(".")
    try:
        if ext == "pdf":
            return _extract_pdf(filepath)
        if ext == "docx":
            return _extract_docx(filepath)
        if ext == "csv":
            return _extract_csv(filepath)
        if ext == "xlsx":
            return _extract_xlsx(filepath)
        if ext == "json":
            return _extract_json(filepath)

        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as exc:
        return f"[Document extraction error: {type(exc).__name__}: {exc}]"


def _extract_pdf(filepath):
    from pypdf import PdfReader
    reader = PdfReader(filepath)
    pages = []
    for number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {number}]\n{text}")
    return "\n\n".join(pages)


def _extract_docx(filepath):
    import docx
    doc = docx.Document(filepath)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table_no, table in enumerate(doc.tables, start=1):
        parts.append(f"[Table {table_no}]")
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_csv(filepath):
    import pandas as pd
    df = pd.read_csv(filepath)
    return df.to_string(index=False)


def _extract_xlsx(filepath):
    import pandas as pd
    sheets = pd.read_excel(filepath, sheet_name=None)
    parts = []
    for name, df in sheets.items():
        parts.append(f"[Sheet: {name}]\n{df.to_string(index=False)}")
    return "\n\n".join(parts)


def _extract_json(filepath):
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        data = json.load(f)
    return json.dumps(data, indent=2, ensure_ascii=False)


# ----------------------------------------------------------------------
# Chunking
# ----------------------------------------------------------------------
def chunk_text(text, chunk_size=None, overlap=None):
    chunk_size = chunk_size or Config.CHUNK_SIZE
    overlap = overlap or Config.CHUNK_OVERLAP
    text = re.sub(r"\r\n?", "\n", (text or "")).strip()
    if not text:
        return []

    # Prefer paragraph boundaries, but keep a predictable character limit.
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    if not paragraphs:
        paragraphs = [text]

    chunks = []
    current = ""
    for paragraph in paragraphs:
        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= chunk_size:
            current = candidate
            continue

        if current:
            chunks.append(current)

        if len(paragraph) <= chunk_size:
            tail = current[-overlap:] if current else ""
            current = f"{tail}\n\n{paragraph}".strip() if tail else paragraph
        else:
            start = 0
            while start < len(paragraph):
                end = min(start + chunk_size, len(paragraph))
                part = paragraph[start:end].strip()
                if part:
                    chunks.append(part)
                if end == len(paragraph):
                    break
                start = max(0, end - overlap)
            current = ""

    if current:
        chunks.append(current)

    return chunks


# ----------------------------------------------------------------------
# Retrieval — hybrid search (TF-IDF + BM25 + keyword) with a lexical
# rerank pass and structured citation metadata.
# ----------------------------------------------------------------------
_TOKEN_RE = re.compile(r"[\w\u0B80-\u0BFF]{2,}")


def _normalize(text):
    return re.sub(r"\s+", " ", (text or "").lower()).strip()


def _tokenize(text):
    return _TOKEN_RE.findall(_normalize(text))


def _keyword_score(query, text):
    q = set(_tokenize(query))
    if not q:
        return 0.0
    t = set(_tokenize(text))
    return len(q & t) / max(1, len(q))


def _minmax(values):
    if not values:
        return values
    lo, hi = min(values), max(values)
    if hi - lo < 1e-9:
        return [0.0 for _ in values]
    return [(v - lo) / (hi - lo) for v in values]


def _tfidf_scores(query, texts):
    """Dense-ish semantic-lexical signal: word n-grams catch meaning/vocabulary
    overlap, char n-grams catch Tamil/Tanglish spelling variation and typos."""
    if not (SKLEARN_OK and len(texts) > 1):
        return [0.0] * len(texts)
    try:
        word = TfidfVectorizer(analyzer="word", ngram_range=(1, 2), sublinear_tf=True, max_features=30000)
        char = TfidfVectorizer(analyzer="char_wb", ngram_range=(3, 5), sublinear_tf=True, max_features=50000)
        wm = word.fit_transform(texts + [query])
        cm = char.fit_transform(texts + [query])
        ws = cosine_similarity(wm[-1], wm[:-1]).ravel()
        cs = cosine_similarity(cm[-1], cm[:-1]).ravel()
        return [0.70 * w + 0.30 * c for w, c in zip(ws, cs)]
    except ValueError:
        return [0.0] * len(texts)


def _bm25_scores(query, texts, k1=1.5, b=0.75):
    """Pure-python Okapi BM25 — a classic sparse/lexical ranking signal that
    complements TF-IDF cosine similarity (which normalises away document
    length and term-frequency saturation differently). No extra dependency
    needed, so this keeps working on any deployment target."""
    q_terms = _tokenize(query)
    if not q_terms:
        return [0.0] * len(texts)
    doc_tokens = [_tokenize(t) for t in texts]
    doc_lens = [len(d) or 1 for d in doc_tokens]
    avgdl = sum(doc_lens) / max(1, len(doc_lens))
    n_docs = len(texts)

    df = {}
    for term in set(q_terms):
        df[term] = sum(1 for d in doc_tokens if term in d)

    idf = {}
    for term, freq in df.items():
        idf[term] = max(0.0, __import__("math").log((n_docs - freq + 0.5) / (freq + 0.5) + 1))

    scores = []
    for tokens, dlen in zip(doc_tokens, doc_lens):
        if not tokens:
            scores.append(0.0)
            continue
        tf = {}
        for tok in tokens:
            tf[tok] = tf.get(tok, 0) + 1
        score = 0.0
        for term in q_terms:
            f = tf.get(term, 0)
            if f == 0:
                continue
            numer = f * (k1 + 1)
            denom = f + k1 * (1 - b + b * dlen / avgdl)
            score += idf.get(term, 0.0) * (numer / denom)
        scores.append(score)
    return scores


def _reciprocal_rank_fusion(*score_lists, k=60):
    """Standard hybrid-search fusion: rank each signal independently, then
    combine by reciprocal rank so no single signal's raw scale dominates."""
    n = len(score_lists[0]) if score_lists else 0
    rrf = [0.0] * n
    for scores in score_lists:
        order = sorted(range(n), key=lambda i: scores[i], reverse=True)
        for rank, idx in enumerate(order):
            if scores[idx] <= 0:
                continue
            rrf[idx] += 1.0 / (k + rank + 1)
    return rrf


def _rerank(query, candidates):
    """Lightweight cross-check reranker applied to the shortlist from hybrid
    retrieval. Rewards exact phrase containment and full query-term coverage
    (signals a single embedding/BM25 score can miss), and mildly penalises
    chunks so short or so long that a match is likely coincidental."""
    q_norm = _normalize(query)
    q_terms = set(_tokenize(query))
    rescored = []
    for row, base_score in candidates:
        text = row["chunk_text"]
        t_norm = _normalize(text)
        t_terms = set(_tokenize(text))
        coverage = len(q_terms & t_terms) / max(1, len(q_terms))
        phrase_bonus = 0.25 if len(q_norm) > 6 and q_norm in t_norm else 0.0
        length_penalty = 0.05 if len(text) < 40 else 0.0
        rerank_signal = coverage + phrase_bonus - length_penalty
        final = 0.6 * base_score + 0.4 * min(1.0, max(0.0, rerank_signal))
        rescored.append((row, final))
    rescored.sort(key=lambda pair: pair[1], reverse=True)
    return rescored


def retrieve_relevant_chunks(query, chunk_rows, top_k=None, return_scores=False):
    """Hybrid retrieval pipeline: TF-IDF + BM25 + keyword-overlap candidates
    fused with reciprocal rank fusion, then reranked with a lexical
    cross-check pass before the final top_k cut."""
    top_k = top_k or Config.TOP_K_CHUNKS
    if not chunk_rows:
        return []

    texts = [c["chunk_text"] for c in chunk_rows]
    tfidf = _tfidf_scores(query, texts)
    bm25 = _bm25_scores(query, texts)
    keyword = [_keyword_score(query, t) for t in texts]

    rrf = _reciprocal_rank_fusion(tfidf, bm25, keyword)
    hybrid = [
        0.55 * a + 0.30 * b + 0.10 * c + 0.05 * (d * 10)
        for a, b, c, d in zip(_minmax(tfidf), _minmax(bm25), keyword, rrf)
    ]

    ranked = sorted(zip(hybrid, chunk_rows), key=lambda pair: pair[0], reverse=True)
    # Do not pretend a random chunk is relevant — cut noise before reranking.
    shortlist = [(row, score) for score, row in ranked if score >= Config.RETRIEVAL_MIN_SCORE][:max(top_k * 5, 15)]
    reranked = _rerank(query, shortlist)
    selected = reranked[:top_k]

    if return_scores:
        return selected
    return [row for row, _score in selected]


def build_citations(selected):
    """Turn retrieved (row, score) pairs into numbered citation metadata the
    UI and the LLM prompt can both reference, e.g. '[1]'."""
    citations = []
    for i, (row, score) in enumerate(selected, start=1):
        text = row["chunk_text"]
        page_match = re.search(r"\[Page (\d+)\]", text)
        snippet = re.sub(r"\[Page \d+\]", "", text).strip()
        snippet = re.sub(r"\s+", " ", snippet)[:220]
        citations.append({
            "index": i,
            "filename": row["filename"],
            "page": int(page_match.group(1)) if page_match else None,
            "snippet": snippet,
            "confidence": round(max(0.0, min(1.0, score)) * 100),
        })
    return citations


# ----------------------------------------------------------------------
# Groq web search
# ----------------------------------------------------------------------
def web_search(query, max_results=5):
    """
    Kept as a compatibility helper. With Groq Compound, web search is
    performed by Groq itself, so no separate Tavily key is required.
    """
    return None, "Groq Compound handles web search automatically."


# ----------------------------------------------------------------------
# LLM
# ----------------------------------------------------------------------
SYSTEM_PROMPT = """You are RAGORA. Answer in the user's Tamil/Tanglish/English style.
Use DOCUMENT EVIDENCE only when it actually supports the answer. Do not invent document facts.
The DOCUMENT EVIDENCE is numbered like [1], [2], [3] in the order given. When a sentence you
write is supported by one of those items, add its bracket number right after it, e.g. "...அப்படி
irukum [1]." Only cite numbers that were actually given to you. Be concise and direct. If the
document does not contain the answer, the caller may provide web evidence instead.
"""

def build_messages(history, user_message, doc_context=None):
    # Keep the entire request intentionally tiny. This is critical for the 8K TPM org limit.
    messages = [{"role":"system","content":SYSTEM_PROMPT}]
    for m in history[-2:]:
        c=(m.get("content") or "").strip()[:180]
        if c:
            messages.append({"role":"assistant" if m.get("role")=="assistant" else "user","content":c})
    if doc_context:
        messages.append({"role":"system","content":"DOCUMENT EVIDENCE:\n"+doc_context[:1800]})
    messages.append({"role":"user","content":(user_message or "")[:700]})
    return messages

def _groq_request(messages, model=None, max_tokens=180, compound=False):
    model = model or (_WEB_MODEL if compound else _resolve_llm_model())
    payload = {"model": model, "messages": messages, "stream": False}
    if not compound:
        payload.update({"temperature":0.2,"max_completion_tokens":max_tokens,"reasoning_effort":"low"})
    headers={"Authorization":f"Bearer {Config.LLM_API_KEY}","Content-Type":"application/json"}
    return requests.post(Config.GROQ_BASE_URL.rstrip("/")+"/chat/completions",headers=headers,json=payload,timeout=Config.LLM_TIMEOUT)

def _error_detail(resp):
    try:
        e=(resp.json().get("error") or {})
        return str(e.get("message") or e.get("type") or "") if isinstance(e,dict) else str(e)
    except Exception:
        return (resp.text or "")[:400]

def _rate_limited(resp):
    return bool(resp is not None and resp.status_code==429)

def _extract_compound_sources(message):
    sources=[]
    for tool in (message.get("executed_tools") or []):
        if not isinstance(tool,dict): continue
        results=tool.get("search_results") or tool.get("output") or []
        if isinstance(results,dict): results=results.get("results") or results.get("items") or [results]
        if not isinstance(results,list): continue
        for item in results:
            if isinstance(item,dict):
                url=item.get("url") or item.get("link"); title=item.get("title") or item.get("name") or url
                if url and str(url).startswith(("http://","https://")): sources.append({"title":str(title)[:120],"url":url})
    seen=set(); out=[]
    for x in sources:
        if x["url"] not in seen: seen.add(x["url"]); out.append(x)
    return out[:6]

def _compound_web_answer(user_message, history):
    messages=[{"role":"user","content":f"Answer this using live web search when needed. Be concise and answer in the user's language.\nQuestion: {(user_message or '')[:700]}"}]
    try:
        resp=_groq_request(messages,_WEB_MODEL,max_tokens=0,compound=True)
        if resp.ok:
            msg=((resp.json().get("choices") or [{}])[0].get("message") or {})
            answer=(msg.get("content") or "").strip()
            if answer: return {"answer":answer,"used_web":True,"sources":_extract_compound_sources(msg)}
        return None
    except requests.RequestException:
        return None

def _duckduckgo_web_context(query):
    try:
        r=requests.get("https://api.duckduckgo.com/",params={"q":query[:400],"format":"json","no_html":1,"skip_disambig":1},timeout=6,headers={"User-Agent":"RAGORA/2.0"})
        if not r.ok: return None,[]
        d=r.json(); items=[]
        if d.get("AbstractText"):
            items.append({"title":d.get("Heading") or "Web result","url":d.get("AbstractURL") or "https://duckduckgo.com/","snippet":d["AbstractText"]})
        for t in (d.get("RelatedTopics") or []):
            if isinstance(t,dict) and t.get("Text"): items.append({"title":t.get("Text","")[:100],"url":t.get("FirstURL") or "https://duckduckgo.com/","snippet":t.get("Text","")})
        items=items[:4]
        if not items:return None,[]
        context="\n\n".join(f"SOURCE: {x['title']}\nSNIPPET: {x['snippet']}" for x in items)
        return context,[{"title":x["title"],"url":x["url"]} for x in items]
    except Exception:return None,[]

def _fallback_document_answer(user_message, doc_context):
    # No-error fallback when the 20B TPM bucket is exhausted. Return the most relevant evidence.
    parts=[p.strip() for p in (doc_context or "").split("\n---\n") if p.strip()]
    if not parts:return "I couldn't generate the AI answer right now, but no document evidence matched this question."
    return "Based on the uploaded document:\n\n"+parts[0][:1200]

def _normal_answer(history,user_message,doc_context=None,web_context=None,web_sources=None):
    messages=build_messages(history,user_message,doc_context)
    if web_context: messages.insert(-1,{"role":"system","content":"WEB EVIDENCE:\n"+web_context[:1800]})
    try:
        resp=_groq_request(messages,_resolve_llm_model(),max_tokens=180,compound=False)
        if resp.ok:
            msg=((resp.json().get("choices") or [{}])[0].get("message") or {})
            answer=(msg.get("content") or "").strip()
            if answer:return {"answer":answer,"used_web":bool(web_context),"sources":web_sources or []}
        # Never retry the same 429. A retry would consume the same org TPM bucket again.
        if _rate_limited(resp) or resp.status_code in (413,500,502,503,504):
            if doc_context:return {"answer":_fallback_document_answer(user_message,doc_context),"used_web":False,"sources":[]}
            if web_context:return {"answer":"Web search found these results, but the AI summary limit is temporarily busy.\n\n"+web_context[:1400],"used_web":True,"sources":web_sources or []}
        return {"answer":"I couldn't generate the answer right now. Please try again in a moment.","used_web":bool(web_context),"sources":web_sources or []}
    except requests.RequestException:
        if doc_context:return {"answer":_fallback_document_answer(user_message,doc_context),"used_web":False,"sources":[]}
        return {"answer":"The AI service is temporarily unavailable. Please try again shortly.","used_web":bool(web_context),"sources":web_sources or []}

def _needs_web_search(user_message):
    text=_normalize(user_message)
    return any(w in text for w in ("latest","today","now","current","recent","news","weather","price","score","schedule","2026","இன்று","இப்போ","தற்போது","நேற்று","நாளை"))

def generate_answer(history,user_message,doc_context=None):
    if not Config.LLM_API_KEY:
        return {"answer":"Groq API key configure pannala. .env-la GROQ_API_KEY add pannunga.","used_web":False,"sources":[]}
    # Current/live questions and weak/no document matches use Compound Mini (70K TPM on free tier).
    if not doc_context or _needs_web_search(user_message):
        web=_compound_web_answer(user_message,history)
        if web:return web
        web_context,sources=_duckduckgo_web_context(user_message)
        if web_context:
            return _normal_answer(history,user_message,None,web_context,sources)
        return _normal_answer(history,user_message,None)
    return _normal_answer(history,user_message,doc_context)

def generate_title(first_message):
    text=(first_message or "").strip().replace("\n"," ")
    if len(text)<=48:return text or "New Chat"
    return text[:48].rsplit(" ",1)[0]+"..."


# ----------------------------------------------------------------------
# AI Chat (friendly companion) — separate personality, separate flow.
# ----------------------------------------------------------------------
COMPANION_SYSTEM_PROMPT = """You are "AI Chat", RAGORA's friendly side companion. You are NOT the
document-RAG assistant — you are just here to chat, like a warm friend checking in.

PERSONALITY
- Talk like a close, caring friend in natural Tanglish/Tamil/English — mirror whatever
  language or mix the user uses.
- Sound casual and human-warm: "Good morning! Eppadi irukeenga today?", "Sapadu aachaa?",
  "Enna panreenga?" — small, genuine check-ins, not a formal assistant greeting.
- Stay upbeat and positive. Gently encourage the user, celebrate small wins, and reassure
  them things will be fine — without dismissing real problems or giving false promises.
- If the user goes quiet, gives one-word replies, or seems bored/low, notice it warmly and
  offer something light, e.g. "Ennada, mounama irukeenga? Oru joke sollatuma, konjam
  sirikalam!" Keep it playful, never pushy or repeated if they say no.
- Keep replies short and conversational (2-4 sentences) like a real chat, not an essay.
  Use emojis sparingly and naturally if it fits the vibe.
- Never diagnose mental health conditions and never claim certainty about the user's
  private feelings — just respond kindly to what they actually say.

SCOPE
- You can also mention how many documents/chunks the user has collected across RAGORA
  when they ask (a DOCUMENT STATS line may be given to you as context) — say it casually,
  e.g. "Ippo unga collection la 3 documents, 42 chunks irukku!"
- You are not doing document Q&A here — if the user asks a serious document/knowledge
  question, gently point them to the main chat: "Idha main chat-la kேlunga, document details
  ellam adhula clear-ah kedaikum!" then keep the tone light.
- Keep everything else about RAGORA (uploads, RAG answers, exports) exactly as-is; you only
  own this casual side conversation.
"""


def build_companion_messages(history, user_message, stats=None):
    messages=[{"role":"system","content":"You are RAGORA's friendly chat companion. Reply briefly in the user's Tamil/Tanglish/English style."}]
    for m in history[-2:]:
        c=(m.get("content") or "").strip()[:180]
        if c: messages.append({"role":"assistant" if m.get("role")=="assistant" else "user","content":c})
    messages.append({"role":"user","content":(user_message or "")[:500]})
    return messages

def generate_companion_answer(history,user_message,stats=None):
    if not Config.LLM_API_KEY:return {"answer":"Groq API key configure pannala."}
    try:
        resp=_groq_request(build_companion_messages(history,user_message,stats),_resolve_llm_model(),max_tokens=140,compound=False)
        if resp.ok:
            msg=((resp.json().get("choices") or [{}])[0].get("message") or {})
            answer=(msg.get("content") or "").strip()
            if answer:return {"answer":answer}
        return {"answer":"Konjam busy ah irukku 🙂  Try again shortly."}
    except Exception:
        return {"answer":"Konjam busy ah irukku 🙂 Try again shortly."}
